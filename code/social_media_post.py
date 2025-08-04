import openai
import re
from pathlib import Path
import json
from typing import Dict, List
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import unicodedata

class SocialMediaPostGeneratorDOCX:
    def __init__(self, api_key: str):
        """
        Inicjalizacja generatora postów na social media z obsługą DOCX
        
        Args:
            api_key: Klucz API do OpenAI
        """
        self.client = openai.OpenAI(api_key=api_key)
        
    def read_transcript(self, file_path: str) -> str:
        """
        Wczytuje transkrypcję z pliku tekstowego
        
        Args:
            file_path: Ścieżka do pliku z transkrypcją
            
        Returns:
            Zawartość transkrypcji jako string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except FileNotFoundError:
            raise FileNotFoundError(f"Nie znaleziono pliku: {file_path}")
        except Exception as e:
            raise Exception(f"Błąd podczas wczytywania pliku: {e}")
    
    def detect_language(self, text: str) -> str:
        """
        Wykrywa język transkrypcji (polski lub angielski)
        
        Args:
            text: Tekst transkrypcji
            
        Returns:
            'pl' dla polskiego, 'en' dla angielskiego
        """
        # Proste wykrywanie języka na podstawie charakterystycznych słów
        polish_words = ['i', 'w', 'na', 'z', 'do', 'się', 'że', 'jak', 'ale', 'to', 'jest', 'dla']
        english_words = ['the', 'and', 'in', 'to', 'of', 'a', 'that', 'it', 'with', 'for', 'as', 'is']
        
        text_lower = text.lower()
        
        polish_count = sum(1 for word in polish_words if f' {word} ' in text_lower)
        english_count = sum(1 for word in english_words if f' {word} ' in text_lower)
        
        return 'pl' if polish_count > english_count else 'en'
    
    def summarize_transcript(self, transcript: str) -> str:
        """
        Podsumowuje długą transkrypcję do kluczowych informacji
        
        Args:
            transcript: Pełna transkrypcja video
            
        Returns:
            Podsumowanie transkrypcji
        """
        try:
            # Jeśli transkrypcja jest krótka, zwróć ją bez zmian
            if len(transcript) < 2000:
                return transcript
            
            # Podsumuj transkrypcję używając GPT
            summary_prompt = f"""
            Podsumuj poniższą transkrypcję video w 2-3 zdaniach, zachowując kluczowe informacje o temacie i głównych zagadnieniach:
            
            {transcript[:3000]}...
            
            Podsumowanie (2-3 zdania):
            """
            
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Jesteś ekspertem w podsumowywaniu treści. Tworzysz zwięzłe, ale informatywne podsumowania."},
                    {"role": "user", "content": summary_prompt}
                ],
                temperature=0.3,
                max_tokens=200
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            # Fallback - zwróć pierwsze 1000 znaków
            return transcript[:1000] + "..."

    def generate_social_media_post(self, transcript: str, detected_language: str) -> Dict[str, str]:
        """
        Generuje posty na social media w obu językach z emotikonami
        
        Args:
            transcript: Transkrypcja video
            detected_language: Wykryty język transkrypcji
            
        Returns:
            Słownik z postami w języku polskim i angielskim
        """
        
        # Podsumuj transkrypcję jeśli jest za długa
        summarized_transcript = self.summarize_transcript(transcript)
        
        # Krótki przykład postu
        example = """
        Przykład:
        [GUIDE] Complete Guide to Structural Analysis!
        Learn step-by-step how to calculate moments of inertia and center of gravity for complex geometric figures. Perfect for engineering students and professionals. [TOOLS][CHART]
        """
        
        # Skrócony prompt
        prompt = f"""
        Stwórz angażujący post na social media na podstawie transkrypcji:
        
        TRANSCRIPT ({detected_language}):
        {summarized_transcript}
        
        WYMAGANIA:
        - Post w języku polskim z emotikonami
        - Tłumaczenie na angielski z emotikonami
        - 2-3 zdania, angażujące
        - Styl podobny do przykładu
        
        {example}
        
        FORMAT JSON:
        {{
            "polish_post": "treść po polsku",
            "english_post": "treść po angielsku",
            "hashtags_pl": ["#hashtag1", "#hashtag2"],
            "hashtags_en": ["#hashtag1", "#hashtag2"],
            "key_topics": ["temat1", "temat2"]
        }}
        """
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Jesteś ekspertem od social media i content marketingu. Tworzysz angażujące posty z emotikonami na podstawie transkrypcji video."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )
            
            # Parsowanie odpowiedzi JSON
            content = response.choices[0].message.content
            
            # Wyciągnięcie JSON z odpowiedzi
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                # Fallback - jeśli nie ma JSON, zwróć prostą strukturę
                return {
                            "polish_post": "[BLAD] Nie udało się wygenerować posta w języku polskim",
        "english_post": "[BLAD] Failed to generate English post",
                    "hashtags_pl": [],
                    "hashtags_en": [],
                    "key_topics": []
                }
                
        except Exception as e:
            raise Exception(f"Błąd podczas generowania posta: {e}")
    
    def remove_emojis(self, text: str) -> str:
        """
        Usuwa wszystkie emoji z tekstu, zachowując tylko zwykłe znaki
        """
        # Regex pattern for emoji detection
        emoji_pattern = re.compile("["
                                   u"\U0001F600-\U0001F64F"  # emoticons
                                   u"\U0001F300-\U0001F5FF"  # symbols & pictographs
                                   u"\U0001F680-\U0001F6FF"  # transport & map symbols
                                   u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
                                   u"\U00002500-\U00002BEF"  # chinese char
                                   u"\U00002702-\U000027B0"
                                   u"\U00002702-\U000027B0"
                                   u"\U000024C2-\U0001F251"
                                   u"\U0001f926-\U0001f937"
                                   u"\U00010000-\U0010ffff"
                                   u"\u2640-\u2642"
                                   u"\u2600-\u2B55"
                                   u"\u200d"
                                   u"\u23cf"
                                   u"\u23e9"
                                   u"\u231a"
                                   u"\ufe0f"  # dingbats
                                   u"\u3030"
                                   "]+", flags=re.UNICODE)
        
        return emoji_pattern.sub(r'', text)

    def add_emoji_to_hashtags(self, hashtags: List[str], language: str) -> List[str]:
        """Dodaje emotikony do hashtagów"""
        emoji_map = {
            'pl': {
                'mechanika': '[GEAR]',
                'belki': '[BUILDING]',
                'siły': '[POWER]',
                'inżynieria': '[TOOLS]',
                'struktury': '[STRUCTURE]',
                'obliczenia': '[CHART]',
                'analiza': '[ANALYSIS]',
                'projektowanie': '[DESIGN]',
                'materiały': '[MATERIALS]',
                'konstrukcje': '[CONSTRUCTION]'
            },
            'en': {
                'mechanics': '[GEAR]',
                'beams': '[BUILDING]',
                'forces': '[POWER]',
                'engineering': '[TOOLS]',
                'structures': '[STRUCTURE]',
                'calculations': '[CHART]',
                'analysis': '[ANALYSIS]',
                'design': '[DESIGN]',
                'materials': '[MATERIALS]',
                'construction': '[CONSTRUCTION]'
            }
        }
        
        emoji_dict = emoji_map.get(language, {})
        enhanced_hashtags = []
        
        for hashtag in hashtags:
            # Usuń # i znajdź odpowiedni emoji
            tag_text = hashtag.replace('#', '').lower()
            emoji = emoji_dict.get(tag_text, '[PIN]')
            enhanced_hashtags.append(f"{emoji} {hashtag}")
        
        return enhanced_hashtags
    
    def save_posts_docx(self, posts: Dict[str, str], output_file: str):
        """
        Zapisuje wygenerowane posty do pliku DOCX z formatowaniem
        
        Args:
            posts: Słownik z postami
            output_file: Ścieżka do pliku wyjściowego
        """
        try:
            # Tworzenie dokumentu
            doc = Document()
            
            # Ustawienie kodowania UTF-8 dla dokumentu
            doc.core_properties.language = 'pl-PL'
            
            # Dodanie tytułu
            title = doc.add_heading('POSTY NA SOCIAL MEDIA', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dodanie daty
            from datetime import datetime
            date_para = doc.add_paragraph(f"Data utworzenia: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Dodanie separatora
            doc.add_paragraph("=" * 80)
            
            # POST PO POLSKU
            doc.add_heading('POST PO POLSKU', level=1)
            
            # Dodanie posta z formatowaniem
            polish_para = doc.add_paragraph()
            polish_run = polish_para.add_run(posts['polish_post'])
            polish_run.font.name = 'Calibri'
            
            # Dodanie hashtagów
            if posts.get('hashtags_pl'):
                hashtags_pl = self.add_emoji_to_hashtags(posts['hashtags_pl'], 'pl')
                hashtags_para = doc.add_paragraph()
                hashtags_label = hashtags_para.add_run("Hashtagi: ")
                hashtags_label.bold = True
                hashtags_label.font.name = 'Calibri'
                hashtags_text = hashtags_para.add_run(" ".join(hashtags_pl))
                hashtags_text.font.name = 'Calibri'
            
            # Dodanie separatora
            doc.add_paragraph("=" * 80)
            
            # POST PO ANGIELSKU
            doc.add_heading('POST PO ANGIELSKU', level=1)
            
            # Dodanie posta z formatowaniem
            english_para = doc.add_paragraph()
            english_run = english_para.add_run(posts['english_post'])
            english_run.font.name = 'Calibri'
            
            # Dodanie hashtagów
            if posts.get('hashtags_en'):
                hashtags_en = self.add_emoji_to_hashtags(posts['hashtags_en'], 'en')
                hashtags_para = doc.add_paragraph()
                hashtags_label = hashtags_para.add_run("Hashtags: ")
                hashtags_label.bold = True
                hashtags_label.font.name = 'Calibri'
                hashtags_text = hashtags_para.add_run(" ".join(hashtags_en))
                hashtags_text.font.name = 'Calibri'
            
            # Dodanie separatora
            doc.add_paragraph("=" * 80)
            
            # KLUCZOWE TEMATY
            doc.add_heading('KLUCZOWE TEMATY', level=1)
            
            if posts.get('key_topics'):
                topics_para = doc.add_paragraph()
                topics_label = topics_para.add_run("Tematy: ")
                topics_label.bold = True
                topics_label.font.name = 'Calibri'
                topics_text = topics_para.add_run(", ".join(posts['key_topics']))
                topics_text.font.name = 'Calibri'
            
            # Dodanie separatora
            doc.add_paragraph("=" * 80)
            
            # INSTRUKCJE KOPIOWANIA
            doc.add_heading('INSTRUKCJE', level=1)
            instructions = doc.add_paragraph()
            instructions_label = instructions.add_run("Jak używać tego dokumentu:\n")
            instructions_label.bold = True
            instructions_label.font.name = 'Calibri'
            
            instructions.add_run("1. Skopiuj odpowiedni post (polski lub angielski)\n").font.name = 'Calibri'
            instructions.add_run("2. Dodaj hashtagi z sekcji poniżej\n").font.name = 'Calibri'
            instructions.add_run("3. Wklej na LinkedIn, Facebook lub inne platformy\n").font.name = 'Calibri'
            instructions.add_run("4. Dostosuj długość jeśli potrzebne\n").font.name = 'Calibri'
            instructions.add_run("5. Dodaj link do video w komentarzach").font.name = 'Calibri'
            
            # Zapisanie dokumentu
            doc.save(output_file)
            
        except Exception as e:
            raise Exception(f"Błąd podczas zapisywania pliku DOCX: {e}")
    
    def save_posts_txt(self, posts: Dict[str, str], output_file: str):
        """
        Zapisuje wygenerowane posty do pliku tekstowego (fallback)
        
        Args:
            posts: Słownik z postami
            output_file: Ścieżka do pliku wyjściowego
        """
        try:
            with open(output_file, 'w', encoding='utf-8') as file:
                file.write("=== POST PO POLSKU ===\n\n")
                file.write(posts['polish_post'])
                file.write(f"\n\nHashtagi: {' '.join(posts.get('hashtags_pl', []))}")
                
                file.write("\n\n" + "="*50 + "\n\n")
                
                file.write("=== POST PO ANGIELSKU ===\n\n")
                file.write(posts['english_post'])
                file.write(f"\n\nHashtags: {' '.join(posts.get('hashtags_en', []))}")
                
                file.write("\n\n" + "="*50 + "\n\n")
                file.write("=== KLUCZOWE TEMATY ===\n\n")
                file.write(", ".join(posts.get('key_topics', [])))
                
        except Exception as e:
            raise Exception(f"Błąd podczas zapisywania pliku: {e}")
    
    def process_transcript(self, input_file: str, output_file: str = None, format_type: str = "docx"):
        """
        Główna funkcja przetwarzająca transkrypcję
        
        Args:
            input_file: Ścieżka do pliku z transkrypcją
            output_file: Ścieżka do pliku wyjściowego (opcjonalne)
            format_type: Format wyjściowy ("docx" lub "txt")
        """
        print(f"[PLIK] Wczytywanie transkrypcji z: {input_file}")
        transcript = self.read_transcript(input_file)
        
        print("[WYKRYWANIE] Wykrywanie języka...")
        language = self.detect_language(transcript)
        print(f"[JEZYK] Wykryty język: {'Polski' if language == 'pl' else 'Angielski'}")
        
        print("[GENEROWANIE] Generowanie postów na social media z emotikonami...")
        posts = self.generate_social_media_post(transcript, language)
        
        # Usuń emoji z wszystkich postów przed wyświetleniem
        posts['polish_post'] = self.remove_emojis(posts['polish_post'])
        posts['english_post'] = self.remove_emojis(posts['english_post'])
        
        # Wyświetlenie wyników
        print("\n" + "="*60)
        print("POST PO POLSKU:")
        print("="*60)
        print(posts['polish_post'])
        
        print("\n" + "="*60)
        print("POST PO ANGIELSKU:")
        print("="*60)
        print(posts['english_post'])
        
        print(f"\n[HASHTAGI] Hashtagi PL: {' '.join(posts.get('hashtags_pl', []))}")
        print(f"[HASHTAGI] Hashtags EN: {' '.join(posts.get('hashtags_en', []))}")
        print(f"[TEMATY] Kluczowe tematy: {', '.join(posts.get('key_topics', []))}")
        
        # Automatyczne generowanie nazwy pliku wyjściowego jeśli nie podano
        if not output_file:
            input_path = Path(input_file)
            if format_type == "docx":
                output_file = input_path.parent / f"{input_path.stem}_post_social_media.docx"
            else:
                output_file = input_path.parent / f"{input_path.stem}_post_social_media.txt"
        
        # Zapisanie do pliku
        if format_type == "docx":
            self.save_posts_docx(posts, output_file)
            print(f"\n[ZAPISANO] Posty zapisane do DOCX: {output_file}")
        else:
            self.save_posts_txt(posts, output_file)
            print(f"\n[ZAPISANO] Posty zapisane do TXT: {output_file}")
        
        return posts

def main():
    parser = argparse.ArgumentParser(description='Generator postów na social media z transkrypcji video (DOCX/TXT)')
    parser.add_argument('input_file', help='Ścieżka do pliku z transkrypcją (WYMAGANE)')
    parser.add_argument('-o', '--output', help='Ścieżka do pliku wyjściowego (opcjonalne)')
    parser.add_argument('-f', '--format', choices=['docx', 'txt'], default='docx', 
                       help='Format wyjściowy (domyślnie: docx)')
    
    args = parser.parse_args()
    
    # Twój klucz API OpenAI - wbudowany w skrypt
    API_KEY = "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"
    
    try:
        # Sprawdź czy plik istnieje
        if not Path(args.input_file).exists():
            print(f"[BLAD] BŁĄD: Nie znaleziono pliku: {args.input_file}")
            print("Upewnij się, że ścieżka do pliku jest poprawna.")
            return
        
        print(f"[START] Przetwarzam plik: {args.input_file}")
        print(f"[FORMAT] Format wyjściowy: {args.format.upper()}")
        
        generator = SocialMediaPostGeneratorDOCX(API_KEY)
        generator.process_transcript(args.input_file, args.output, args.format)
        
        print("\n[SUKCES] Gotowe! Posty zostały wygenerowane z emotikonami!")
        
    except Exception as e:
        print(f"[BLAD] Błąd: {e}")

if __name__ == "__main__":
    main() 